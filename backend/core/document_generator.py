"""Document Generator - Generate DOCX and Markdown files"""
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from loguru import logger

from config import settings


class DocumentGenerator:
    """Document Generator for DOCX and Markdown"""

    def _get_unique_filename(self, output_dir: Path, base_filename: str) -> str:
        """
        Get unique filename by adding (1), (2), etc. if file already exists

        Args:
            output_dir: Output directory
            base_filename: Base filename with extension

        Returns:
            Unique filename
        """
        output_path = output_dir / base_filename

        # 如果文件不存在，直接返回原文件名
        if not output_path.exists():
            return base_filename

        # 分离文件名和扩展名
        name_part = base_filename.rsplit('.', 1)[0]
        ext_part = '.' + base_filename.rsplit('.', 1)[1] if '.' in base_filename else ''

        # 尝试添加数字后缀
        counter = 1
        while True:
            new_filename = f"{name_part}({counter}){ext_part}"
            new_path = output_dir / new_filename
            if not new_path.exists():
                return new_filename
            counter += 1

            # 防止无限循环，最多尝试1000次
            if counter > 1000:
                import time
                timestamp = int(time.time())
                return f"{name_part}_{timestamp}{ext_part}"

    def _clean_filename(self, filename: str) -> str:
        """
        Clean filename by removing illegal characters

        Args:
            filename: Original filename

        Returns:
            Cleaned filename
        """
        import re
        # 移除或替换非法字符
        cleaned = re.sub(r'[/\\:*?"<>|]', '_', filename)
        # 移除多余的空格和下划线
        cleaned = re.sub(r'[_\s]+', '_', cleaned).strip('_')
        # 限制长度
        if len(cleaned) > 100:
            cleaned = cleaned[:100]
        return cleaned

    def generate_docx(
        self,
        content: str,
        title: Optional[str] = None,
        output_filename: Optional[str] = None,
        output_dir: Optional[str] = None
    ) -> Optional[Path]:
        """
        Generate DOCX document from text content

        Args:
            content: Text content
            title: Document title
            output_filename: Output filename (auto-generate if None)
            output_dir: Custom output directory (use default if None)

        Returns:
            Output file path if successful, None otherwise
        """
        try:
            if not content:
                logger.error("Empty content provided")
                return None

            # Create document
            doc = Document()

            # Add title if provided
            if title:
                heading = doc.add_heading(title, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add timestamp
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            doc.add_paragraph(f"Generated at: {timestamp}")
            doc.add_paragraph()  # Empty line

            # Add content
            # Split content by paragraphs
            paragraphs = content.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    paragraph = doc.add_paragraph(para_text.strip())
                    # Set font
                    for run in paragraph.runs:
                        run.font.size = Pt(12)
                        run.font.name = 'Arial'

            # Generate output filename if not provided
            if not output_filename:
                # 尝试从内容的第一行提取标题作为文件名
                first_line = content.split('\n')[0].strip() if content else ""
                # 移除开头的Markdown标题符号（#）
                title_text = first_line.replace('#', '').strip()
                # 清理文件名中的非法字符
                clean_title = self._clean_filename(title_text)

                if clean_title and len(clean_title) > 0:
                    # DOCX文件名只取前10个字符
                    short_title = clean_title[:10] if len(clean_title) > 10 else clean_title
                    output_filename = f"{short_title}.docx"
                else:
                    # 如果第一行为空或太短，使用时间戳
                    timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
                    output_filename = f"document_{timestamp_str}.docx"

            # Ensure .docx extension
            if not output_filename.endswith('.docx'):
                output_filename += '.docx'

            # Determine output directory
            if output_dir:
                output_base_dir = Path(output_dir)
            else:
                output_base_dir = settings.output_dir

            # Create directory if it doesn't exist
            output_base_dir.mkdir(parents=True, exist_ok=True)

            # Get unique filename to avoid overwriting existing files
            unique_filename = self._get_unique_filename(output_base_dir, output_filename)

            # Save to output directory
            output_path = output_base_dir / unique_filename
            doc.save(str(output_path))

            logger.info(f"DOCX document generated: {output_path}")
            return output_path

        except Exception as e:
            logger.exception(f"Error generating DOCX: {e}")
            return None

    def generate_markdown(
        self,
        content: str,
        title: Optional[str] = None,
        output_filename: Optional[str] = None,
        output_dir: Optional[str] = None
    ) -> Optional[Path]:
        """
        Generate Markdown file from text content

        Args:
            content: Text content
            title: Document title
            output_filename: Output filename (auto-generate if None)
            output_dir: Custom output directory (use default if None)

        Returns:
            Output file path if successful, None otherwise
        """
        try:
            if not content:
                logger.error("Empty content provided")
                return None

            # Build markdown content
            md_parts = []

            # Add title if provided
            if title:
                md_parts.append(f"# {title}\n")

            # Add timestamp
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            md_parts.append(f"*Generated at: {timestamp}*\n")
            md_parts.append("---\n")

            # Add content
            md_parts.append(content)

            md_content = "\n".join(md_parts)

            # Generate output filename if not provided
            if not output_filename:
                # 尝试从内容的第一行提取标题作为文件名
                first_line = content.split('\n')[0].strip() if content else ""
                # 移除开头的Markdown标题符号（#）
                title_text = first_line.replace('#', '').strip()
                # 清理文件名中的非法字符
                clean_title = self._clean_filename(title_text)

                if clean_title and len(clean_title) > 0:
                    output_filename = f"{clean_title}.md"
                else:
                    # 如果第一行为空或太短，使用时间戳
                    timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
                    output_filename = f"document_{timestamp_str}.md"

            # Ensure .md extension
            if not output_filename.endswith('.md'):
                output_filename += '.md'

            # Determine output directory
            if output_dir:
                output_base_dir = Path(output_dir)
            else:
                output_base_dir = settings.output_dir

            # Create directory if it doesn't exist
            output_base_dir.mkdir(parents=True, exist_ok=True)

            # Get unique filename to avoid overwriting existing files
            unique_filename = self._get_unique_filename(output_base_dir, output_filename)

            # Save to output directory
            output_path = output_base_dir / unique_filename
            output_path.write_text(md_content, encoding='utf-8')

            logger.info(f"Markdown file generated: {output_path}")
            return output_path

        except Exception as e:
            logger.exception(f"Error generating Markdown: {e}")
            return None


# Global generator instance
document_generator = DocumentGenerator()
